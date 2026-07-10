#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《材料科学前沿对RISC-V微架构设计的启示与思考》Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
#  创建文档
# ============================================================
doc = Document()

# ============================================================
#  页面设置
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ============================================================
#  样式定义
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
# 中文字体回退
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题样式重新定义
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = 'Times New Roman'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif i == 2:
        heading_style.font.size = Pt(15)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    else:
        heading_style.font.size = Pt(13)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_para(text, bold=False, font_size=12, alignment=None, first_line_indent=True,
             font_name=None, color=None):
    """添加段落的便捷函数"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 两个中文字符
    if alignment is not None:
        p.alignment = alignment

    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name or 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_bold_para(text, font_size=12):
    """添加加粗段落（用于强调句）"""
    return add_para(text, bold=True, font_size=font_size)


def add_image_placeholder(caption_text, fig_num, width_inches=5.5):
    """添加图片占位框 + 图题"""
    # 用一个带边框的文本框作为占位符
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    # 添加占位说明文字
    run = p.add_run(f'〔 图{fig_num}：请在此处插入图片 —— {caption_text} 〕')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True

    # 图题
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_after = Pt(12)
    run_cap = cap.add_run(f'图{fig_num}  {caption_text}')
    run_cap.font.size = Pt(10)
    run_cap.font.name = 'Times New Roman'
    run_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_cap.bold = True


# ============================================================
#  封面
# ============================================================
# 空行
for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5

# 主标题
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
run = p.add_run('跨越物理边界：')
run.font.size = Pt(26)
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
run = p.add_run('材料科学前沿对RISC-V微架构设计的启示与思考')
run.font.size = Pt(26)
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

# 空行
for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5

# 作者信息
info_lines = [
    ('易烽鑫', 16),
    ('北京航空航天大学', 14),
    ('电子与信息工程学院', 14),
    ('2026年6月', 14),
]
for text, size in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.8
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
#  分页 - 摘要
# ============================================================
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(18)
run = p.add_run('摘  要')
run.font.size = Pt(18)
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

abstract_text = (
    '本文是一篇基于《材料科学与工程学科综合课》系列前沿讲座的课程反思报告。'
    '作为一名长期从事RISC-V微架构设计的研究生，本文结合赵立东老师的热电转换材料、'
    '赵士腾老师的高熵合金、肖文龙老师的轻质高强金属、侯慧龙/赵新青/徐惠彬老师团队的'
    '形状记忆合金、茹毅研究员的高温结构材料、李岩老师的生物医用材料以及康鹏老师的AI赋能'
    '材料设计等讲座内容，系统探讨了材料科学前沿成果对处理器微架构设计的跨学科启示。'
    '本文从热管理、封装可靠性、极低功耗、边缘计算架构以及AI辅助设计空间探索等多个维度，'
    '提出了将材料物理特性纳入微架构设计闭环的创新思路，旨在探索"后摩尔时代"下'
    '"材料—架构—算法"三位一体的协同设计新范式。'
)
add_para(abstract_text, first_line_indent=True)

# 关键词
kw = doc.add_paragraph()
kw.paragraph_format.line_spacing = 1.5
kw.paragraph_format.space_before = Pt(12)
run_label = kw.add_run('关键词：')
run_label.font.size = Pt(12)
run_label.bold = True
run_label.font.name = 'Times New Roman'
run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run_value = kw.add_run('RISC-V微架构；材料科学；热电转换；高熵合金；形状记忆合金；AI辅助设计；跨学科协同')
run_value.font.size = Pt(12)
run_value.font.name = 'Times New Roman'
run_value._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
#  分页 - 正文开始
# ============================================================
doc.add_page_break()

# ============================================================
#  一、引言
# ============================================================
doc.add_heading('一、引言', level=1)

add_para(
    '作为一名长期沉浸在数字电路与计算机体系结构领域的研究生，我的日常研究几乎完全被Verilog代码、'
    'RTL仿真波形、时序约束以及性能计数器所填满。在接触《材料科学与工程学科综合课》之前，'
    '我潜意识里一直认为"材料"是物理层面、工艺层面的问题，而"微架构"是逻辑层、算法层的问题，'
    '两者虽有联系，但在我的日常工作中存在一条隐形的鸿沟。'
)

add_para(
    '然而，通过本学期北航材料学院各位知名教授带来的系列前沿讲座，我深刻意识到：在摩尔定律放缓、'
    '芯片步入"后摩尔时代"的今天，微架构设计者如果继续只盯着门级电路和流水线，而忽视底层材料物理'
    '特性的约束与赋能，将无法真正解决高性能芯片面临的"功耗墙"、"散热墙"和"可靠性墙"。特别是对于'
    '我们正在研究的RISC-V核，其开放性和可扩展性不仅提供了架构上的自由，也赋予了我们在物理层面探索'
    '"材料—架构"协同设计的巨大潜力。'
)

add_para(
    '这堂课不仅让我大开眼界，更让我对自己正在设计的RISC-V核的未来方向产生了强烈的反思和重构冲动。'
    '以下，我将结合每位老师的讲座内容，谈谈这些材料前沿技术对我微架构研究方向的具体触动与启示。'
)

# ============================================================
#  二、热电材料
# ============================================================
doc.add_heading('二、从"被动散热"到"主动能量回收"：热电材料对处理器热管理架构的冲击', level=1)

add_para(
    '赵立东老师主讲的《新型高效热电转换材料》是我印象最深的一讲。长期以来，我在设计RISC-V核'
    '（尤其是乱序执行的多发射核）时，最头疼的问题就是高负载下处理器的局部温度过高。传统的解决方法'
    '无非是降低频率、插入门控时钟、或者在物理设计时大量铺陈金属层用于导热，本质上都是在"忍受"热量'
    '的产生，然后试图用风冷或液冷将其"驱逐"出去。这种被动散热机制，在面对数亿乃至数十亿晶体管密度'
    '的现代芯片时，已经暴露出极大的物理瓶颈，也就是我们常说的"热墙"。'
)

add_image_placeholder('热电转换材料与片上热管理架构的协同设计概念示意图', 1)

add_para(
    '赵老师的研究却给了我一个全新的视角：',
    first_line_indent=True
)
add_bold_para('为什么不把废热看作是资源，而是去主动利用它？')

add_para(
    '赵老师介绍的通过能带工程和声子工程提升热电优值（ZT值）的方法，让我联想到在RISC-V芯片中集成'
    '微型热电薄膜的可行性。在我的微架构设计流程中，我通常只评估"动态功耗"和"漏电流功耗"，却从未将'
    '芯片表面的"温差"视作一个设计变量。听完讲座后，我立刻在自己的微架构性能模型里加入了一个设想：'
    '假设我们在RISC-V的整数运算单元（ALU）和浮点运算单元（FPU）下方集成高ZT热电材料，当核心全速'
    '运转产生巨量热量时，上下表面形成的温度梯度可以转化为电能，用来给片上的传感器或低功耗唤醒模块'
    '供电。这不仅仅是物理层面的构想，它对我微架构层的"温度感知调度策略"提出了直接挑战。'
)

add_para(
    '目前我的RISC-V多核调度算法主要考虑任务负载均衡，但在未来，我可能需要设计一种能够主动制造温差'
    '梯度的调度策略——将高强度的计算任务集中在某一核心，让它"发热"，从而驱动热电模块高效发电，这显然'
    '打破了传统微架构设计师追求"温度均衡"的认知，让我认识到微架构与材料工艺之间可以实现"共生"而非'
    '单纯的"妥协"。这种跨学科的认知，对我下一步在微架构中加入"热—电耦合模型"有着深远的指导意义。'
)

# ============================================================
#  三、高熵合金
# ============================================================
doc.add_heading('三、高熵合金与轻质高强金属：对先进封装下RISC-V核可靠性的再思考', level=1)

add_para(
    '赵士腾老师主讲的《高熵合金》与肖文龙老师主讲的《轻质高强金属结构材料》，让我对芯片的"物理可靠性"'
    '有了全新的理解。目前我设计的RISC-V核主要采用标准CMOS工艺流片，在微架构仿真阶段，我只关心时序收敛'
    '（Setup/Hold）和功能正确性。但随着先进封装技术（如Chiplet、3D IC）的普及，片上系统（SoC）不再是'
    '一个单一的硅片，而是由多个Die通过硅通孔（TSV）和微凸点（Micro-bump）堆叠而成。这种三维立体集成'
    '在带来高带宽、低延迟优势的同时，也引入了极其复杂的热机械应力问题。'
)

add_image_placeholder('先进封装（Chiplet/3D IC）中的热机械应力与互联可靠性挑战', 2)

add_para(
    '赵老师提到的高熵合金具备迟滞扩散效应和高硬度特点，这让我意识到，传统互连材料（铜、铝）在高温下的'
    '电迁移失效，在微架构层面其实对应着"关键路径延迟抖动"和"逻辑翻转错误"。如果我们在微架构设计阶段没有'
    '引入材料老化模型，那么仿真验证中通过的优秀性能，可能在实际物理芯片寿命末期根本不成立。肖老师提到的'
    '轻质高强金属在航空航天中的应用，同样映射到3D芯片封装中——随着堆叠层数增加，不同材料热膨胀系数（CTE）'
    '不匹配带来的热机械应力会导致芯片翘曲和断裂。'
)

add_para(
    '这让我反思，在未来的RISC-V微架构研究里，我必须将"物理层应力"纳入微架构级的故障容忍机制。例如，'
    '我是否可以设计一种"应力感知的片上网络（NoC）路由算法"？在芯片运行过程中，结合嵌入式传感器监测到的'
    '应力分布，绕过那些可能因为材料疲劳而存在高阻值的互联通路，从而在逻辑层上延长芯片的整体寿命。这种'
    '"微架构—材料力学"的交叉，确实是我之前完全没考虑过的方向，也是我未来在验证RISC-V核长期可靠性时'
    '必须引入的新维度。'
)

# ============================================================
#  四、形状记忆合金
# ============================================================
doc.add_heading('四、形状记忆合金与高温结构材料：物理开关带来的极低功耗新思路', level=1)

add_para(
    '侯慧龙、赵新青、徐惠彬老师团队的《基于固态相变的形状记忆合金基础前沿与空天应用》，让我对芯片的'
    '"电源管理"产生了颠覆性的想法。在RISC-V微架构中，为了降低待机功耗，我们常使用"时钟门控"和"电源门控"，'
    '但电路层面的电源开关始终存在漏电流问题。无论我们在微架构层面把电源门控设计得多精细，晶体管自身的'
    '物理泄漏是无法彻底消除的。'
)

add_image_placeholder('形状记忆合金（SMA）固态相变机理与微机械继电器概念', 3)

add_para(
    '形状记忆合金（SMA）在特定温度下会因马氏体相变而产生巨大的形变和回复应力，这给了我一个大胆的设想：',
    first_line_indent=True
)
add_bold_para('能否在芯片内部利用形状记忆合金制造微机械继电器？')
add_para(
    '虽然SMA的响应速度在毫秒级，比晶体管慢得多，但它可以作为"粗粒度电源域开关"。比如，在我的RISC-V架构中，'
    '有一个长期闲置的AI协处理器，传统方式下即使关断电源，漏电流依然存在。如果采用SMA热机械开关，完全物理'
    '断开电源连接，理论上可以实现接近绝对零功耗的待机。对于电池供电的边缘计算设备而言，这无疑是一个极具'
    '吸引力的方案。这对于我在设计面向低功耗物联网场景的RISC-V核时，提供了一个极具前瞻性的物理实现选项。'
)

add_para(
    '此外，茹毅研究员主讲的《发动机用高温材料复杂工程问题思考》也给了我启发。虽然航空发动机涡轮叶片与芯片'
    '在尺度上相差悬殊，但关于高温下的热障涂层和热机械疲劳的物理本质是完全相通的。芯片在工作时，热循环引起'
    '的焊点疲劳和层间应力开裂，本质上与高温合金的失效机理相似。这提醒我，在评估RISC-V处理器的长期可靠性时，'
    '不能只看电磁兼容，更要借鉴材料科学中的热疲劳模型，来重新定义微架构设计的"安全温度阈值"。茹老师提到的'
    '"复杂工程问题"概念，让我意识到把芯片散热看作一个单纯的工程数学问题是不够的，它本质上是一个受材料热力学'
    '制约的复杂物理问题。'
)

# ============================================================
#  五、生物医用材料
# ============================================================
doc.add_heading('五、先进生物医用材料与植入式计算：边缘架构的新需求', level=1)

add_para(
    '李岩老师的《先进生物医用材料与医疗器械》讲座，让我看到了RISC-V微架构在特定垂直领域落地的巨大潜力。'
    '我们课题组在研究RISC-V核时，往往以通用计算性能为目标，追求更高的IPC和更复杂的乱序执行逻辑。然而，'
    '李老师提到的植入式医疗器械对芯片有"超低功耗"、"高生物相容性"、"无线供电与数据传输"等苛刻要求，'
    '这让我深刻反思通用架构在极端场景下的局限性。'
)

add_image_placeholder('植入式医疗设备中的芯片架构需求：生物相容封装与超低功耗约束', 4)

add_para(
    '这让我意识到，',
    first_line_indent=True
)
add_bold_para('微架构设计不能脱离应用场景的材料约束。')
add_para(
    '如果我的RISC-V核要植入人体，那么封装材料必须经过特殊的生物相容性处理，这意味着芯片的顶部金属层'
    '（Top Metal）必须专门预留出大面积的天线接口和能量采集接口，而这些接口在传统的标准单元库设计中通常是'
    '视为"干扰源"去回避的。此外，为了防止植入物发热损伤组织，我的微架构必须引入极其激进的"功耗封顶"策略'
    '——一旦检测到某功能单元产生局部热点，立即强制降频甚至通过指令预取中断来降低活跃度。这要求我在流水线'
    '设计中，提前预留更多的"应急降级"通路，而非仅仅追求峰值性能。李老师的讲座让我明白，一个优秀的RISC-V'
    '微架构师，不应该只是通晓指令集和流水线，更应该懂得如何根据生物相容材料的物理特性来反向定制数据通路，'
    '这样才能真正把芯片做成救死扶伤的医疗器械核心。'
)

# ============================================================
#  六、AI赋能设计
# ============================================================
doc.add_heading('六、人工智能赋能设计：是"工具"更是"方法论"的迁移', level=1)

add_para(
    '康鹏老师的《人工智能赋能材料设计研发概述》可能是与我这名微架构设计师最同频共振的一节课。他详细介绍'
    '了如何利用机器学习、图神经网络等AI技术加速新材料的发现和性能预测，将传统以"试错法"为主的材料研发转变'
    '为"数据驱动"的范式。康老师提到的材料基因组计划和机器学习势能函数，让我脑海中立刻浮现出我们设计芯片时'
    '的"设计空间探索（DSE）"问题。'
)

add_image_placeholder('从"AI for Materials"到"AI for Microarchitecture"的范式迁移', 5)

add_para(
    '这正是我们目前RISC-V微架构设计最需要的"降维打击"。当前的RISC-V核设计空间极其庞大：流水线深度'
    '（5级、7级、10级）、保留站数量、重排序缓存（ROB）大小、分支预测器配置、缓存关联度等参数组合可达数百万种，'
    '这被称为"组合爆炸"。用仿真器去跑一遍所有组合是不现实的，而过去我们依靠工程师经验去拍脑袋定参数，往往错过'
    '最优解。康老师的讲座让我坚定了一个想法：我完全可以将"AI for Materials"的范式迁移到'
    '"AI for Microarchitecture"中。'
)

add_para(
    '我计划在自己的研究中，把RISC-V微架构的参数空间视作"合金成分空间"，把仿真的功耗、面积、性能（PPA）视作'
    '"材料性能指标"。通过构建一个深度强化学习（DRL）代理，让它在有限的仿真次数下，自动探索出最适合特定应用场景'
    '的微架构参数组合。这种"架构自动寻优"的思路，本质上与AI辅助材料设计寻找最佳成分比例是完全一致的，甚至在'
    '损失函数的设计上都可以相互借鉴。这让我摆脱了对传统手工调参的依赖，为我未来的微架构研究注入了全新的方法论，'
    '也让我深刻认识到，软件算法（AI）和硬件材料在更高维度上，其实遵循着相似的优化逻辑。'
)

# ============================================================
#  七、综合反思
# ============================================================
doc.add_heading('七、综合反思与自身研究方向的再定位', level=1)

add_para(
    '回顾这一系列课程，我最大的收获在于',
    first_line_indent=True
)
add_bold_para('打破了"微架构隔离"的认知茧房。')
add_para(
    '以前，我只在门级和寄存器传输级（RTL）工作，认为CPU就是一堆逻辑门和寄存器的排列组合；现在我的视野'
    '扩展到了热力学、力学和电磁学层面，开始明白CPU也是一个需要呼吸、会发热、会衰老的物理实体。基于这些启发，'
    '我对自己的RISC-V微架构设计研究做出了如下深度的调整和长远规划：'
)

# 6个小节
subsections = [
    ('1. 建立"热—力—电"耦合的早期评估模型',
     '在传统的Verilog RTL仿真之外，我需要增加一个物理层降阶模型。利用讲座中获得的材料物理参数'
     '（如热导率、CTE、热电Seebeck系数），初步评估不同微架构配置下的片上温度分布和热应力。具体来说，'
     '我不会等到后端物理设计做完才去评估热，而是在架构设计阶段，就调用一个用Python编写的"材料物理代理模型"，'
     '把缓存访问率和执行单元活动率折算成单位面积的热流密度，从而在流片之前，我就知道这个核在特定封装材料下'
     '是否会因为过热而发生时序违例或崩溃。'),

    ('2. 设计"材料感知"的动态功耗管理单元（PMU）',
     '传统的PMU只是根据负载切换电压和频率，这是一种开环的控制。我计划在我的RISC-V核中，尝试设计一种闭环的'
     '"热力耦合"PMU。不再简单地根据负载切换，而是在PMU中引入温度变化率和温差梯度的监测。如果检测到局部温差'
     '过大，恰好进入某种热电薄膜材料的高效工作区间，那么我可以通过微架构指令调度，适度增加该区域的负荷以回收'
     '热量，把废热变成电能送给L2缓存供电；反之，如果温差过小，则启动全局散热。这种基于材料物理特性的动态反馈'
     '机制，是传统微架构文献中极少提及的创新点。'),

    ('3. 在核心微架构中预留"异构材料接口"',
     '在做RTL编码和芯片顶层规划时，我现在会有意识地预留物理接口。比如在芯片的顶层金属布线规划中，我会为未来'
     '的热电薄膜和形状记忆合金开关预留足够的物理布线通道和特殊焊盘（Pad），确保微架构层面的逻辑控制信号可以'
     '驱动这些新型材料器件。这意味着，我的RTL代码中不仅要规定A端口接B端口，还要考虑高电流驱动能力，以便未来'
     '在物理实现时能与微机电系统（MEMS）或热电模块兼容。'),

    ('4. 拥抱AI辅助架构探索',
     '我已经开始着手构建一个小型的"RISC-V微架构参数空间探索数据集"，尝试用图神经网络替代传统的仿真验证。在RTL'
     '综合前就得到比较可靠的PPA预估，能够将我的设计迭代周期从数周缩短到数天。此外，我也在思考是否可以用迁移学习'
     '的方法，把材料科学领域的AI预训练模型参数，迁移到微架构的功耗预测上，因为两者在特征提取（如周期性、局部性）'
     '上存在数学上的相似性。'),

    ('5. 引入多物理场故障注入验证',
     '我过去做验证，主要是通过编写随机指令测试用例，关注数据流的正确性。这门课启发我，真正的"鲁棒设计"必须能抵抗'
     '热机械应力。我计划在未来的验证流程中，设计一种"热机械故障注入"方法——基于材料物理模型，模拟在特定温度循环下，'
     '某些连线因为电迁移导致电阻值变大，从而产生"延迟故障"。通过观察我的RISC-V微架构在这种模拟故障下是否能够通过'
     'ECC纠错或指令重试来恢复，来衡量其可靠性，这彻底颠覆了我原有的验证框架。'),

    ('6. 建立"材料—架构—算法"三位一体的理论框架',
     '通过这门课，我深刻领悟到，最优的芯片设计往往是三者协同的结果。我在写微架构论文时，以前只提逻辑优化，现在'
     '我会尝试引入"物理感知"的评价指标。比如，在评估一个分支预测器改进方案时，不仅要看它提升了多少IPC，还要计算'
     '它因为降低了误预测率，从而减少了流水线冲刷带来的功耗浪涌，进而降低了芯片的平均温度，延长了材料寿命。我希望'
     '在未来，能够把我所设计的RISC-V核，构建成一个能够主动适应封装材料物理属性变化的自适应平台。'),
]

for title, text in subsections:
    doc.add_heading(title, level=2)
    add_para(text)

# ============================================================
#  八、结语
# ============================================================
doc.add_heading('八、结语', level=1)

add_para(
    '《材料科学与工程学科综合课》对我而言，不仅仅是一次知识的拓展，更是一次思维范式的冲击与重塑。'
    '从李宜彬老师的多功能多尺度复合材料，到赵立东老师的热电转换，再到康鹏老师的AI赋能，各位教授展示的'
    '不仅仅是材料学科的高精尖成果，更是一种"从原子到系统"的全局观。'
)

add_para(
    '作为一名RISC-V微架构设计者，我深刻认识到，未来的高性能计算芯片，其核心竞争力将不仅取决于指令集的优劣'
    '或流水线的效率，更取决于设计者能否在物理材料的约束下寻找最优解，以及能否利用新兴材料的特性去重塑计算架构。'
    '以前我们常说"架构决定了芯片的上限"，现在我明白，是',
    first_line_indent=True
)
add_bold_para('"材料决定了芯片的物理边界，而架构决定了我们能在多大程度上逼近这个边界"。')
add_para(
    '通过这门课程，我学会了用更宽广、更跨学科的眼光去看待自己手中的那几万行Verilog代码，也让我对如何设计出'
    '真正面向未来、适应复杂物理环境的高能效RISC-V处理器有了更清晰的路径。'
)

add_para(
    '在未来的科研道路上，我将积极践行这门课程带给我的跨学科交叉理念，努力在自己的微架构研究领域，主动拥抱'
    '材料科学、热力学与人工智能，成为一个"懂物理、懂架构、懂算法"的复合型研究人员。我期待在不久的将来，能够'
    '把我从课堂上学到的这些材料前沿知识，真正化为我芯片设计中那些富有远见的小改进，这，就是这门综合课给我'
    '最宝贵、最深远的财富。'
)

# ============================================================
#  保存
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), '课程报告_材料科学与RISC-V微架构.docx')
doc.save(output_path)
print(f'文档已生成：{output_path}')
